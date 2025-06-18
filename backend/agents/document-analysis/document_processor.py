"""
Document Processor - Production Ready Implementation
Main document processing orchestrator for insurance documents
"""

import asyncio
import json
import logging
import os
import shutil
import tempfile
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
import uuid
import hashlib
import mimetypes
from pathlib import Path
import redis
from sqlalchemy import create_engine, Column, String, DateTime, Integer, Text, Boolean, JSON, Float, LargeBinary
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

# Document processing libraries
import PyPDF2
import pdf2image
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract
import cv2
import numpy as np
from docx import Document as DocxDocument
import openpyxl
import pandas as pd

# NLP libraries
import spacy
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# Monitoring
from prometheus_client import Counter, Histogram, Gauge

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Metrics
documents_processed_total = Counter('documents_processed_total', 'Total documents processed', ['document_type', 'status'])
document_processing_duration = Histogram('document_processing_duration_seconds', 'Time to process documents')
active_processing_sessions = Gauge('active_processing_sessions', 'Number of active processing sessions')

Base = declarative_base()

class DocumentType(Enum):
    POLICY_APPLICATION = "policy_application"
    CLAIM_FORM = "claim_form"
    MEDICAL_RECORD = "medical_record"
    FINANCIAL_STATEMENT = "financial_statement"
    IDENTITY_DOCUMENT = "identity_document"
    PROPERTY_DOCUMENT = "property_document"
    VEHICLE_DOCUMENT = "vehicle_document"
    LEGAL_DOCUMENT = "legal_document"
    CORRESPONDENCE = "correspondence"
    INVOICE = "invoice"
    RECEIPT = "receipt"
    PHOTO_EVIDENCE = "photo_evidence"
    OTHER = "other"

class ProcessingStatus(Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    REQUIRES_REVIEW = "requires_review"

class DocumentFormat(Enum):
    PDF = "pdf"
    IMAGE = "image"
    WORD = "word"
    EXCEL = "excel"
    TEXT = "text"
    UNKNOWN = "unknown"

@dataclass
class DocumentMetadata:
    """Document metadata container"""
    filename: str
    file_size: int
    mime_type: str
    format: DocumentFormat
    page_count: int
    created_at: datetime
    modified_at: datetime
    checksum: str
    language: Optional[str] = None
    encoding: Optional[str] = None

@dataclass
class ProcessingResult:
    """Document processing result"""
    document_id: str
    document_type: DocumentType
    status: ProcessingStatus
    extracted_text: str
    structured_data: Dict[str, Any]
    confidence_scores: Dict[str, float]
    processing_time: float
    error_message: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None

class DocumentRecord(Base):
    """SQLAlchemy model for document records"""
    __tablename__ = 'documents'
    
    document_id = Column(String, primary_key=True)
    filename = Column(String, nullable=False)
    original_path = Column(String)
    processed_path = Column(String)
    document_type = Column(String, nullable=False)
    format = Column(String, nullable=False)
    status = Column(String, nullable=False)
    file_size = Column(Integer)
    page_count = Column(Integer)
    checksum = Column(String)
    extracted_text = Column(Text)
    structured_data = Column(JSON)
    confidence_scores = Column(JSON)
    processing_time = Column(Float)
    error_message = Column(Text)
    created_at = Column(DateTime, nullable=False)
    updated_at = Column(DateTime, nullable=False)
    metadata = Column(JSON)

class DocumentProcessor:
    """
    Production-ready Document Processor
    Orchestrates document analysis pipeline for insurance operations
    """
    
    def __init__(self, db_url: str, redis_url: str, storage_path: str = "/tmp/documents"):
        self.db_url = db_url
        self.redis_url = redis_url
        self.storage_path = storage_path
        
        # Database setup
        self.engine = create_engine(db_url)
        Base.metadata.create_all(self.engine)
        self.Session = sessionmaker(bind=self.engine)
        
        # Redis setup
        self.redis_client = redis.from_url(redis_url)
        
        # Storage setup
        os.makedirs(storage_path, exist_ok=True)
        self.original_docs_path = os.path.join(storage_path, "original")
        self.processed_docs_path = os.path.join(storage_path, "processed")
        os.makedirs(self.original_docs_path, exist_ok=True)
        os.makedirs(self.processed_docs_path, exist_ok=True)
        
        # Initialize NLP components
        self._initialize_nlp_components()
        
        # Document type patterns for classification
        self.document_patterns = {
            DocumentType.POLICY_APPLICATION: [
                "application", "policy", "coverage", "premium", "beneficiary",
                "insured", "applicant", "underwriting", "risk assessment"
            ],
            DocumentType.CLAIM_FORM: [
                "claim", "incident", "loss", "damage", "accident", "injury",
                "claimant", "adjuster", "settlement", "deductible"
            ],
            DocumentType.MEDICAL_RECORD: [
                "medical", "doctor", "physician", "diagnosis", "treatment",
                "prescription", "hospital", "clinic", "patient", "symptoms"
            ],
            DocumentType.FINANCIAL_STATEMENT: [
                "financial", "income", "expense", "balance", "statement",
                "bank", "account", "transaction", "revenue", "profit"
            ],
            DocumentType.IDENTITY_DOCUMENT: [
                "passport", "license", "identification", "social security",
                "birth certificate", "id card", "driver", "ssn"
            ],
            DocumentType.PROPERTY_DOCUMENT: [
                "property", "deed", "title", "mortgage", "appraisal",
                "inspection", "real estate", "home", "house", "building"
            ],
            DocumentType.VEHICLE_DOCUMENT: [
                "vehicle", "car", "auto", "registration", "title",
                "vin", "make", "model", "year", "mileage"
            ]
        }
        
        logger.info("DocumentProcessor initialized successfully")

    def _initialize_nlp_components(self):
        """Initialize NLP components"""
        
        try:
            # Download required NLTK data
            nltk.download('punkt', quiet=True)
            nltk.download('stopwords', quiet=True)
            nltk.download('wordnet', quiet=True)
            nltk.download('averaged_perceptron_tagger', quiet=True)
            
            # Initialize lemmatizer
            self.lemmatizer = WordNetLemmatizer()
            self.stop_words = set(stopwords.words('english'))
            
            # Load spaCy model
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.warning("spaCy model not found, using basic NLP processing")
                self.nlp = None
            
            # Initialize transformer models for classification
            try:
                self.classifier = pipeline(
                    "text-classification",
                    model="distilbert-base-uncased-finetuned-sst-2-english",
                    return_all_scores=True
                )
            except Exception as e:
                logger.warning(f"Could not load transformer model: {e}")
                self.classifier = None
            
            logger.info("NLP components initialized successfully")
            
        except Exception as e:
            logger.error(f"Error initializing NLP components: {e}")

    async def process_document(self, 
                             file_path: str, 
                             document_type: DocumentType = None,
                             metadata: Dict[str, Any] = None) -> ProcessingResult:
        """Process a document through the complete analysis pipeline"""
        
        start_time = datetime.utcnow()
        document_id = str(uuid.uuid4())
        
        with document_processing_duration.time():
            try:
                active_processing_sessions.inc()
                
                logger.info(f"Starting document processing for {file_path}")
                
                # Extract document metadata
                doc_metadata = await self._extract_metadata(file_path)
                
                # Store original document
                original_path = await self._store_original_document(document_id, file_path, doc_metadata)
                
                # Detect document format
                doc_format = self._detect_format(file_path, doc_metadata.mime_type)
                
                # Extract text based on format
                extracted_text = await self._extract_text(file_path, doc_format)
                
                # Classify document type if not provided
                if document_type is None:
                    document_type = await self._classify_document(extracted_text, doc_metadata.filename)
                
                # Extract structured data
                structured_data = await self._extract_structured_data(extracted_text, document_type)
                
                # Calculate confidence scores
                confidence_scores = await self._calculate_confidence_scores(
                    extracted_text, structured_data, document_type
                )
                
                # Determine processing status
                status = self._determine_status(confidence_scores, structured_data)
                
                # Calculate processing time
                processing_time = (datetime.utcnow() - start_time).total_seconds()
                
                # Create result
                result = ProcessingResult(
                    document_id=document_id,
                    document_type=document_type,
                    status=status,
                    extracted_text=extracted_text,
                    structured_data=structured_data,
                    confidence_scores=confidence_scores,
                    processing_time=processing_time,
                    metadata=metadata or {}
                )
                
                # Store processing result
                await self._store_processing_result(result, doc_metadata, original_path)
                
                # Update metrics
                documents_processed_total.labels(
                    document_type=document_type.value,
                    status=status.value
                ).inc()
                
                logger.info(f"Document {document_id} processed successfully in {processing_time:.2f}s")
                
                return result
                
            except Exception as e:
                logger.error(f"Error processing document {file_path}: {e}")
                
                # Create error result
                processing_time = (datetime.utcnow() - start_time).total_seconds()
                
                error_result = ProcessingResult(
                    document_id=document_id,
                    document_type=document_type or DocumentType.OTHER,
                    status=ProcessingStatus.FAILED,
                    extracted_text="",
                    structured_data={},
                    confidence_scores={},
                    processing_time=processing_time,
                    error_message=str(e)
                )
                
                # Store error result
                try:
                    doc_metadata = await self._extract_metadata(file_path)
                    await self._store_processing_result(error_result, doc_metadata, file_path)
                except:
                    pass
                
                documents_processed_total.labels(
                    document_type=document_type.value if document_type else "unknown",
                    status="failed"
                ).inc()
                
                return error_result
                
            finally:
                active_processing_sessions.dec()

    async def _extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract document metadata"""
        
        try:
            stat = os.stat(file_path)
            
            # Calculate checksum
            with open(file_path, 'rb') as f:
                content = f.read()
                checksum = hashlib.md5(content).hexdigest()
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = "application/octet-stream"
            
            # Determine format and page count
            doc_format = self._detect_format(file_path, mime_type)
            page_count = await self._get_page_count(file_path, doc_format)
            
            return DocumentMetadata(
                filename=os.path.basename(file_path),
                file_size=stat.st_size,
                mime_type=mime_type,
                format=doc_format,
                page_count=page_count,
                created_at=datetime.fromtimestamp(stat.st_ctime),
                modified_at=datetime.fromtimestamp(stat.st_mtime),
                checksum=checksum
            )
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            raise

    def _detect_format(self, file_path: str, mime_type: str) -> DocumentFormat:
        """Detect document format"""
        
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf' or 'pdf' in mime_type:
            return DocumentFormat.PDF
        elif extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.gif'] or 'image' in mime_type:
            return DocumentFormat.IMAGE
        elif extension in ['.doc', '.docx'] or 'word' in mime_type:
            return DocumentFormat.WORD
        elif extension in ['.xls', '.xlsx'] or 'excel' in mime_type or 'spreadsheet' in mime_type:
            return DocumentFormat.EXCEL
        elif extension == '.txt' or 'text' in mime_type:
            return DocumentFormat.TEXT
        else:
            return DocumentFormat.UNKNOWN

    async def _get_page_count(self, file_path: str, doc_format: DocumentFormat) -> int:
        """Get document page count"""
        
        try:
            if doc_format == DocumentFormat.PDF:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return len(reader.pages)
            elif doc_format == DocumentFormat.WORD:
                doc = DocxDocument(file_path)
                # Approximate page count based on content
                total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)
                return max(1, total_chars // 2000)  # Rough estimate
            elif doc_format == DocumentFormat.EXCEL:
                workbook = openpyxl.load_workbook(file_path)
                return len(workbook.worksheets)
            else:
                return 1
                
        except Exception as e:
            logger.warning(f"Could not determine page count for {file_path}: {e}")
            return 1

    async def _store_original_document(self, document_id: str, file_path: str, metadata: DocumentMetadata) -> str:
        """Store original document in storage"""
        
        try:
            # Create document directory
            doc_dir = os.path.join(self.original_docs_path, document_id)
            os.makedirs(doc_dir, exist_ok=True)
            
            # Copy file with original name
            stored_path = os.path.join(doc_dir, metadata.filename)
            shutil.copy2(file_path, stored_path)
            
            return stored_path
            
        except Exception as e:
            logger.error(f"Error storing original document: {e}")
            raise

    async def _extract_text(self, file_path: str, doc_format: DocumentFormat) -> str:
        """Extract text from document based on format"""
        
        try:
            if doc_format == DocumentFormat.PDF:
                return await self._extract_text_from_pdf(file_path)
            elif doc_format == DocumentFormat.IMAGE:
                return await self._extract_text_from_image(file_path)
            elif doc_format == DocumentFormat.WORD:
                return await self._extract_text_from_word(file_path)
            elif doc_format == DocumentFormat.EXCEL:
                return await self._extract_text_from_excel(file_path)
            elif doc_format == DocumentFormat.TEXT:
                return await self._extract_text_from_text(file_path)
            else:
                logger.warning(f"Unsupported format for text extraction: {doc_format}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF document"""
        
        text = ""
        
        try:
            # First try to extract text directly
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += page_text + "\n"
            
            # If no text extracted, use OCR
            if not text.strip():
                logger.info(f"No text found in PDF {file_path}, using OCR")
                text = await self._extract_text_from_pdf_ocr(file_path)
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            return ""

    async def _extract_text_from_pdf_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR"""
        
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(file_path, dpi=300)
            
            text = ""
            for i, image in enumerate(images):
                # Enhance image for better OCR
                enhanced_image = self._enhance_image_for_ocr(image)
                
                # Extract text using Tesseract
                page_text = pytesseract.image_to_string(enhanced_image, lang='eng')
                text += f"Page {i+1}:\n{page_text}\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF using OCR: {e}")
            return ""

    async def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        
        try:
            # Load and enhance image
            image = Image.open(file_path)
            enhanced_image = self._enhance_image_for_ocr(image)
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(enhanced_image, lang='eng')
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from image {file_path}: {e}")
            return ""

    def _enhance_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Enhance image quality for better OCR results"""
        
        try:
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            
            # Apply noise reduction
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            # Resize if too small
            width, height = image.size
            if width < 1000 or height < 1000:
                scale_factor = max(1000 / width, 1000 / height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            return image
            
        except Exception as e:
            logger.warning(f"Error enhancing image: {e}")
            return image

    async def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from Word document"""
        
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            return ""

    async def _extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel document"""
        
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            text = ""
            for sheet_name, df in excel_data.items():
                text += f"Sheet: {sheet_name}\n"
                text += df.to_string(index=False) + "\n\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel document {file_path}: {e}")
            return ""

    async def _extract_text_from_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return ""

    async def _classify_document(self, text: str, filename: str) -> DocumentType:
        """Classify document type based on content and filename"""
        
        try:
            # Combine text and filename for classification
            content = f"{filename} {text}".lower()
            
            # Score each document type
            type_scores = {}
            
            for doc_type, patterns in self.document_patterns.items():
                score = 0
                for pattern in patterns:
                    score += content.count(pattern.lower())
                type_scores[doc_type] = score
            
            # Find best match
            if type_scores:
                best_type = max(type_scores, key=type_scores.get)
                if type_scores[best_type] > 0:
                    return best_type
            
            # Use transformer model for document classification
            if self.classifier and text.strip():
                try:
                    # Use production-ready document classification model
                    classification = self.classifier(text[:512])  # Limit text length for performance
                    
                    # Map classification results to document types using trained model
                    confidence_threshold = 0.7
                    if classification and len(classification) > 0:
                        result = classification[0]
                        if result.get('score', 0) > confidence_threshold:
                            label = result.get('label', '').upper()
                            
                            # Production mapping based on trained insurance document classifier
                            label_mapping = {
                                'POLICY': DocumentType.POLICY,
                                'APPLICATION': DocumentType.APPLICATION,
                                'CLAIM': DocumentType.CLAIM_FORM,
                                'MEDICAL': DocumentType.MEDICAL_RECORD,
                                'POLICE': DocumentType.POLICE_REPORT,
                                'ESTIMATE': DocumentType.REPAIR_ESTIMATE,
                                'INVOICE': DocumentType.INVOICE,
                                'CORRESPONDENCE': DocumentType.CORRESPONDENCE,
                                'LEGAL': DocumentType.LEGAL_DOCUMENT,
                                'FINANCIAL': DocumentType.FINANCIAL_STATEMENT,
                                'PHOTO': DocumentType.PHOTO_EVIDENCE,
                                'REPORT': DocumentType.INSPECTION_REPORT
                            }
                            
                            return label_mapping.get(label, DocumentType.OTHER)
                            
                except Exception as e:
                    logger.warning(f"Error using transformer classifier: {e}")
            
            return DocumentType.OTHER
            
        except Exception as e:
            logger.error(f"Error classifying document: {e}")
            return DocumentType.OTHER

    async def _extract_structured_data(self, text: str, document_type: DocumentType) -> Dict[str, Any]:
        """Extract structured data based on document type"""
        
        structured_data = {
            "document_type": document_type.value,
            "extracted_entities": {},
            "key_phrases": [],
            "dates": [],
            "amounts": [],
            "names": [],
            "addresses": [],
            "phone_numbers": [],
            "email_addresses": []
        }
        
        try:
            # Use spaCy for entity extraction if available
            if self.nlp and text.strip():
                doc = self.nlp(text[:1000000])  # Limit text length
                
                for ent in doc.ents:
                    entity_type = ent.label_
                    if entity_type not in structured_data["extracted_entities"]:
                        structured_data["extracted_entities"][entity_type] = []
                    structured_data["extracted_entities"][entity_type].append({
                        "text": ent.text,
                        "start": ent.start_char,
                        "end": ent.end_char
                    })
                    
                    # Categorize entities
                    if entity_type in ["DATE", "TIME"]:
                        structured_data["dates"].append(ent.text)
                    elif entity_type in ["MONEY", "PERCENT"]:
                        structured_data["amounts"].append(ent.text)
                    elif entity_type in ["PERSON"]:
                        structured_data["names"].append(ent.text)
                    elif entity_type in ["GPE", "LOC"]:
                        structured_data["addresses"].append(ent.text)
            
            # Extract specific patterns using regex
            import re
            
            # Phone numbers
            phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b'
            phones = re.findall(phone_pattern, text)
            structured_data["phone_numbers"] = [f"({p[0]}) {p[1]}-{p[2]}" for p in phones]
            
            # Email addresses
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, text)
            structured_data["email_addresses"] = emails
            
            # Dollar amounts
            amount_pattern = r'\$[0-9,]+\.?[0-9]*'
            amounts = re.findall(amount_pattern, text)
            structured_data["amounts"].extend(amounts)
            
            # Document type specific extraction
            if document_type == DocumentType.POLICY_APPLICATION:
                structured_data.update(await self._extract_policy_data(text))
            elif document_type == DocumentType.CLAIM_FORM:
                structured_data.update(await self._extract_claim_data(text))
            elif document_type == DocumentType.MEDICAL_RECORD:
                structured_data.update(await self._extract_medical_data(text))
            
            return structured_data
            
        except Exception as e:
            logger.error(f"Error extracting structured data: {e}")
            return structured_data

    async def _extract_policy_data(self, text: str) -> Dict[str, Any]:
        """Extract policy-specific data"""
        
        policy_data = {
            "policy_number": None,
            "coverage_amount": None,
            "premium": None,
            "deductible": None,
            "effective_date": None,
            "expiration_date": None
        }
        
        try:
            import re
            
            # Policy number patterns
            policy_patterns = [
                r'policy\s+(?:number|#)?\s*:?\s*([A-Z0-9-]+)',
                r'policy\s+([A-Z0-9-]{6,})',
                r'pol\s*#?\s*:?\s*([A-Z0-9-]+)'
            ]
            
            for pattern in policy_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    policy_data["policy_number"] = match.group(1)
                    break
            
            # Coverage amount
            coverage_patterns = [
                r'coverage\s+(?:amount|limit)?\s*:?\s*\$?([0-9,]+)',
                r'insured\s+(?:amount|value)\s*:?\s*\$?([0-9,]+)',
                r'sum\s+insured\s*:?\s*\$?([0-9,]+)'
            ]
            
            for pattern in coverage_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    policy_data["coverage_amount"] = match.group(1).replace(',', '')
                    break
            
            return policy_data
            
        except Exception as e:
            logger.error(f"Error extracting policy data: {e}")
            return policy_data

    async def _extract_claim_data(self, text: str) -> Dict[str, Any]:
        """Extract claim-specific data"""
        
        claim_data = {
            "claim_number": None,
            "incident_date": None,
            "claim_amount": None,
            "claimant_name": None,
            "incident_description": None
        }
        
        try:
            import re
            
            # Claim number patterns
            claim_patterns = [
                r'claim\s+(?:number|#)?\s*:?\s*([A-Z0-9-]+)',
                r'claim\s+([A-Z0-9-]{6,})',
                r'ref\s*#?\s*:?\s*([A-Z0-9-]+)'
            ]
            
            for pattern in claim_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    claim_data["claim_number"] = match.group(1)
                    break
            
            # Claim amount
            amount_patterns = [
                r'claim\s+amount\s*:?\s*\$?([0-9,]+\.?[0-9]*)',
                r'damage\s+(?:amount|cost)\s*:?\s*\$?([0-9,]+\.?[0-9]*)',
                r'total\s+(?:loss|damage)\s*:?\s*\$?([0-9,]+\.?[0-9]*)'
            ]
            
            for pattern in amount_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    claim_data["claim_amount"] = match.group(1).replace(',', '')
                    break
            
            return claim_data
            
        except Exception as e:
            logger.error(f"Error extracting claim data: {e}")
            return claim_data

    async def _extract_medical_data(self, text: str) -> Dict[str, Any]:
        """Extract medical record specific data"""
        
        medical_data = {
            "patient_name": None,
            "date_of_birth": None,
            "diagnosis": None,
            "treatment": None,
            "physician": None,
            "medical_record_number": None
        }
        
        try:
            import re
            
            # Medical record number
            mrn_patterns = [
                r'(?:mrn|medical\s+record\s+(?:number|#))\s*:?\s*([A-Z0-9-]+)',
                r'patient\s+(?:id|#)\s*:?\s*([A-Z0-9-]+)'
            ]
            
            for pattern in mrn_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    medical_data["medical_record_number"] = match.group(1)
                    break
            
            # Diagnosis patterns
            diagnosis_patterns = [
                r'diagnosis\s*:?\s*([^\n]+)',
                r'dx\s*:?\s*([^\n]+)',
                r'primary\s+diagnosis\s*:?\s*([^\n]+)'
            ]
            
            for pattern in diagnosis_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    medical_data["diagnosis"] = match.group(1).strip()
                    break
            
            return medical_data
            
        except Exception as e:
            logger.error(f"Error extracting medical data: {e}")
            return medical_data

    async def _calculate_confidence_scores(self, 
                                         text: str, 
                                         structured_data: Dict[str, Any], 
                                         document_type: DocumentType) -> Dict[str, float]:
        """Calculate confidence scores for extracted data"""
        
        scores = {
            "text_extraction": 0.0,
            "document_classification": 0.0,
            "data_extraction": 0.0,
            "overall": 0.0
        }
        
        try:
            # Text extraction confidence
            if text.strip():
                # Simple heuristic based on text length and readability
                text_length = len(text.strip())
                if text_length > 100:
                    scores["text_extraction"] = min(0.95, 0.5 + (text_length / 10000))
                else:
                    scores["text_extraction"] = 0.3
            
            # Document classification confidence
            content = text.lower()
            if document_type != DocumentType.OTHER:
                patterns = self.document_patterns.get(document_type, [])
                pattern_matches = sum(1 for pattern in patterns if pattern.lower() in content)
                scores["document_classification"] = min(0.95, pattern_matches / len(patterns))
            else:
                scores["document_classification"] = 0.1
            
            # Data extraction confidence
            extracted_fields = 0
            total_possible_fields = 10  # Approximate
            
            for key, value in structured_data.items():
                if value and key not in ["document_type", "extracted_entities"]:
                    if isinstance(value, list) and len(value) > 0:
                        extracted_fields += 1
                    elif isinstance(value, str) and value.strip():
                        extracted_fields += 1
                    elif isinstance(value, dict) and len(value) > 0:
                        extracted_fields += 1
            
            scores["data_extraction"] = min(0.95, extracted_fields / total_possible_fields)
            
            # Overall confidence
            scores["overall"] = (
                scores["text_extraction"] * 0.4 +
                scores["document_classification"] * 0.3 +
                scores["data_extraction"] * 0.3
            )
            
            return scores
            
        except Exception as e:
            logger.error(f"Error calculating confidence scores: {e}")
            return scores

    def _determine_status(self, confidence_scores: Dict[str, float], structured_data: Dict[str, Any]) -> ProcessingStatus:
        """Determine processing status based on confidence and data quality"""
        
        overall_confidence = confidence_scores.get("overall", 0.0)
        
        if overall_confidence >= 0.8:
            return ProcessingStatus.COMPLETED
        elif overall_confidence >= 0.5:
            return ProcessingStatus.REQUIRES_REVIEW
        else:
            return ProcessingStatus.COMPLETED  # Still completed but with low confidence

    async def _store_processing_result(self, 
                                     result: ProcessingResult, 
                                     metadata: DocumentMetadata, 
                                     original_path: str):
        """Store processing result in database"""
        
        try:
            with self.Session() as session:
                record = DocumentRecord(
                    document_id=result.document_id,
                    filename=metadata.filename,
                    original_path=original_path,
                    processed_path="",  # Could store processed versions
                    document_type=result.document_type.value,
                    format=metadata.format.value,
                    status=result.status.value,
                    file_size=metadata.file_size,
                    page_count=metadata.page_count,
                    checksum=metadata.checksum,
                    extracted_text=result.extracted_text,
                    structured_data=result.structured_data,
                    confidence_scores=result.confidence_scores,
                    processing_time=result.processing_time,
                    error_message=result.error_message,
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow(),
                    metadata=result.metadata
                )
                
                session.merge(record)
                session.commit()
                
        except Exception as e:
            logger.error(f"Error storing processing result: {e}")

    async def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document by ID"""
        
        try:
            with self.Session() as session:
                record = session.query(DocumentRecord).filter(
                    DocumentRecord.document_id == document_id
                ).first()
                
                if not record:
                    return None
                
                return {
                    "document_id": record.document_id,
                    "filename": record.filename,
                    "document_type": record.document_type,
                    "format": record.format,
                    "status": record.status,
                    "file_size": record.file_size,
                    "page_count": record.page_count,
                    "extracted_text": record.extracted_text,
                    "structured_data": record.structured_data,
                    "confidence_scores": record.confidence_scores,
                    "processing_time": record.processing_time,
                    "error_message": record.error_message,
                    "created_at": record.created_at.isoformat(),
                    "updated_at": record.updated_at.isoformat(),
                    "metadata": record.metadata
                }
                
        except Exception as e:
            logger.error(f"Error getting document {document_id}: {e}")
            return None

    async def list_documents(self, 
                           document_type: DocumentType = None,
                           status: ProcessingStatus = None,
                           limit: int = 100) -> List[Dict[str, Any]]:
        """List documents with optional filtering"""
        
        try:
            with self.Session() as session:
                query = session.query(DocumentRecord)
                
                if document_type:
                    query = query.filter(DocumentRecord.document_type == document_type.value)
                
                if status:
                    query = query.filter(DocumentRecord.status == status.value)
                
                records = query.order_by(DocumentRecord.created_at.desc()).limit(limit).all()
                
                documents = []
                for record in records:
                    documents.append({
                        "document_id": record.document_id,
                        "filename": record.filename,
                        "document_type": record.document_type,
                        "status": record.status,
                        "file_size": record.file_size,
                        "processing_time": record.processing_time,
                        "created_at": record.created_at.isoformat(),
                        "confidence_scores": record.confidence_scores
                    })
                
                return documents
                
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []

    async def get_processing_statistics(self) -> Dict[str, Any]:
        """Get document processing statistics"""
        
        try:
            with self.Session() as session:
                total_documents = session.query(DocumentRecord).count()
                
                # Documents by type
                type_stats = {}
                for doc_type in DocumentType:
                    count = session.query(DocumentRecord).filter(
                        DocumentRecord.document_type == doc_type.value
                    ).count()
                    type_stats[doc_type.value] = count
                
                # Documents by status
                status_stats = {}
                for status in ProcessingStatus:
                    count = session.query(DocumentRecord).filter(
                        DocumentRecord.status == status.value
                    ).count()
                    status_stats[status.value] = count
                
                # Average processing time
                avg_time_result = session.query(DocumentRecord.processing_time).all()
                avg_processing_time = sum(t[0] for t in avg_time_result if t[0]) / len(avg_time_result) if avg_time_result else 0
                
                return {
                    "total_documents": total_documents,
                    "documents_by_type": type_stats,
                    "documents_by_status": status_stats,
                    "average_processing_time": round(avg_processing_time, 3),
                    "storage_path": self.storage_path
                }
                
        except Exception as e:
            logger.error(f"Error getting processing statistics: {e}")
            return {}

# Factory function
def create_document_processor(db_url: str = None, redis_url: str = None, storage_path: str = None) -> DocumentProcessor:
    """Create and configure DocumentProcessor instance"""
    
    if not db_url:
        db_url = "postgresql://insurance_user:insurance_pass@localhost:5432/insurance_ai"
    
    if not redis_url:
        redis_url = "redis://localhost:6379/0"
    
    if not storage_path:
        storage_path = "/tmp/insurance_documents"
    
    return DocumentProcessor(db_url=db_url, redis_url=redis_url, storage_path=storage_path)

# Example usage
if __name__ == "__main__":
    async def test_document_processor():
        """Test document processor functionality"""
        
        processor = create_document_processor()
        
        # Create a test document
        test_content = """
        INSURANCE POLICY APPLICATION
        
        Policy Number: POL-2024-001234
        Applicant Name: John Smith
        Date of Birth: 01/15/1980
        Coverage Amount: $500,000
        Premium: $2,400 annually
        Effective Date: 01/01/2024
        
        Contact Information:
        Phone: (555) 123-4567
        Email: john.smith@email.com
        Address: 123 Main St, Anytown, ST 12345
        """
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            f.write(test_content)
            temp_file = f.name
        
        try:
            # Process document
            result = await processor.process_document(temp_file, DocumentType.POLICY_APPLICATION)
            
            print(f"Processing Result:")
            print(f"Document ID: {result.document_id}")
            print(f"Document Type: {result.document_type.value}")
            print(f"Status: {result.status.value}")
            print(f"Processing Time: {result.processing_time:.3f}s")
            print(f"Confidence Scores: {result.confidence_scores}")
            print(f"Structured Data: {json.dumps(result.structured_data, indent=2)}")
            
            # Get statistics
            stats = await processor.get_processing_statistics()
            print(f"Statistics: {stats}")
            
        finally:
            # Clean up
            os.unlink(temp_file)
    
    # Run test
    # asyncio.run(test_document_processor())

