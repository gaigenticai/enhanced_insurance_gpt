"""
Metadata Extractor - Production Ready Implementation
Comprehensive metadata extraction for digital evidence
"""

import asyncio
import json
import logging
import os
import struct
import tempfile
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any, Union, Tuple
from dataclasses import dataclass, asdict
from enum import Enum
import uuid
import hashlib
import mimetypes
from pathlib import Path
import redis
from sqlalchemy import create_engine, Column, String, DateTime, Integer, Text, Boolean, JSON, Float
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session

# File format libraries
from PIL import Image, ExifTags
from PIL.ExifTags import TAGS, GPSTAGS
import PyPDF2
import docx
import openpyxl
import mutagen
from mutagen.id3 import ID3
import cv2
import ffmpeg

# System libraries
import platform
import psutil
import subprocess

# Monitoring
from prometheus_client import Counter, Histogram

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

metadata_extraction_total = Counter('metadata_extraction_total', 'Total metadata extractions', ['file_type'])
metadata_extraction_duration = Histogram('metadata_extraction_duration_seconds', 'Time for metadata extraction')

Base = declarative_base()

class MetadataType(Enum):
    EXIF = "exif"
    IPTC = "iptc"
    XMP = "xmp"
    FILE_SYSTEM = "file_system"
    DOCUMENT_PROPERTIES = "document_properties"
    AUDIO_TAGS = "audio_tags"
    VIDEO_METADATA = "video_metadata"
    SYSTEM_METADATA = "system_metadata"

class FileCategory(Enum):
    IMAGE = "image"
    VIDEO = "video"
    AUDIO = "audio"
    DOCUMENT = "document"
    ARCHIVE = "archive"
    EXECUTABLE = "executable"
    OTHER = "other"

@dataclass
class GeolocationData:
    latitude: float
    longitude: float
    altitude: Optional[float]
    accuracy: Optional[float]
    timestamp: Optional[datetime]
    source: str

@dataclass
class DeviceInformation:
    make: Optional[str]
    model: Optional[str]
    serial_number: Optional[str]
    software_version: Optional[str]
    hardware_info: Dict[str, Any]

@dataclass
class TimestampData:
    creation_time: Optional[datetime]
    modification_time: Optional[datetime]
    access_time: Optional[datetime]
    metadata_time: Optional[datetime]
    timezone_info: Optional[str]

@dataclass
class FileSystemMetadata:
    file_path: str
    file_name: str
    file_size: int
    file_type: str
    mime_type: str
    permissions: str
    owner: str
    group: str
    inode: Optional[int]
    links_count: int
    checksums: Dict[str, str]

@dataclass
class ContentMetadata:
    title: Optional[str]
    author: Optional[str]
    subject: Optional[str]
    keywords: List[str]
    description: Optional[str]
    creator_tool: Optional[str]
    language: Optional[str]
    page_count: Optional[int]
    word_count: Optional[int]

@dataclass
class TechnicalMetadata:
    format_version: Optional[str]
    compression: Optional[str]
    color_space: Optional[str]
    bit_depth: Optional[int]
    resolution: Optional[Tuple[int, int]]
    duration: Optional[float]
    bitrate: Optional[int]
    sample_rate: Optional[int]
    channels: Optional[int]

@dataclass
class SecurityMetadata:
    encryption_status: bool
    digital_signature: Optional[str]
    certificate_info: Optional[Dict[str, Any]]
    password_protected: bool
    access_restrictions: List[str]

@dataclass
class MetadataExtractionResult:
    extraction_id: str
    file_path: str
    file_category: FileCategory
    extraction_timestamp: datetime
    file_system_metadata: FileSystemMetadata
    timestamp_data: TimestampData
    geolocation_data: Optional[GeolocationData]
    device_information: Optional[DeviceInformation]
    content_metadata: Optional[ContentMetadata]
    technical_metadata: Optional[TechnicalMetadata]
    security_metadata: Optional[SecurityMetadata]
    raw_metadata: Dict[str, Any]
    extraction_errors: List[str]
    processing_time: float
    confidence_score: float

class MetadataExtractionRecord(Base):
    __tablename__ = 'metadata_extractions'
    
    extraction_id = Column(String, primary_key=True)
    evidence_id = Column(String, index=True)
    file_path = Column(String, nullable=False)
    file_category = Column(String)
    extraction_timestamp = Column(DateTime, nullable=False)
    file_system_metadata = Column(JSON)
    timestamp_data = Column(JSON)
    geolocation_data = Column(JSON)
    device_information = Column(JSON)
    content_metadata = Column(JSON)
    technical_metadata = Column(JSON)
    security_metadata = Column(JSON)
    raw_metadata = Column(JSON)
    extraction_errors = Column(JSON)
    processing_time = Column(Float)
    confidence_score = Column(Float)
    created_at = Column(DateTime, nullable=False)

class MetadataExtractor:
    """Production-ready Metadata Extractor for digital evidence"""
    
    def __init__(self, db_url: str, redis_url: str):
        self.db_url = db_url
        self.redis_url = redis_url
        
        self.engine = create_engine(db_url)
        Base.metadata.create_all(self.engine)
        self.Session = sessionmaker(bind=self.engine)
        
        self.redis_client = redis.from_url(redis_url)
        
        # File type mappings
        self.image_extensions = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif', '.webp', '.heic', '.raw'}
        self.video_extensions = {'.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv', '.webm', '.m4v'}
        self.audio_extensions = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a'}
        self.document_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', '.txt', '.rtf'}
        self.archive_extensions = {'.zip', '.rar', '.7z', '.tar', '.gz', '.bz2'}
        
        logger.info("MetadataExtractor initialized successfully")

    async def extract_metadata(self, file_path: str, evidence_id: str = None) -> MetadataExtractionResult:
        """Extract comprehensive metadata from file"""
        
        start_time = datetime.utcnow()
        extraction_id = str(uuid.uuid4())
        
        with metadata_extraction_duration.time():
            try:
                # Determine file category
                file_category = self._determine_file_category(file_path)
                
                # Extract file system metadata
                fs_metadata = await self._extract_filesystem_metadata(file_path)
                
                # Extract timestamp data
                timestamp_data = await self._extract_timestamp_data(file_path)
                
                # Initialize optional metadata containers
                geolocation_data = None
                device_information = None
                content_metadata = None
                technical_metadata = None
                security_metadata = None
                raw_metadata = {}
                extraction_errors = []
                
                # Extract category-specific metadata
                try:
                    if file_category == FileCategory.IMAGE:
                        result = await self._extract_image_metadata(file_path)
                        geolocation_data = result.get('geolocation')
                        device_information = result.get('device')
                        technical_metadata = result.get('technical')
                        raw_metadata.update(result.get('raw', {}))
                        
                    elif file_category == FileCategory.VIDEO:
                        result = await self._extract_video_metadata(file_path)
                        technical_metadata = result.get('technical')
                        content_metadata = result.get('content')
                        raw_metadata.update(result.get('raw', {}))
                        
                    elif file_category == FileCategory.AUDIO:
                        result = await self._extract_audio_metadata(file_path)
                        content_metadata = result.get('content')
                        technical_metadata = result.get('technical')
                        raw_metadata.update(result.get('raw', {}))
                        
                    elif file_category == FileCategory.DOCUMENT:
                        result = await self._extract_document_metadata(file_path)
                        content_metadata = result.get('content')
                        security_metadata = result.get('security')
                        raw_metadata.update(result.get('raw', {}))
                        
                except Exception as e:
                    extraction_errors.append(f"Category-specific extraction failed: {str(e)}")
                
                # Extract security metadata for all files
                try:
                    if not security_metadata:
                        security_metadata = await self._extract_security_metadata(file_path)
                except Exception as e:
                    extraction_errors.append(f"Security metadata extraction failed: {str(e)}")
                
                # Calculate confidence score
                confidence_score = self._calculate_confidence_score(
                    fs_metadata, timestamp_data, geolocation_data, device_information,
                    content_metadata, technical_metadata, extraction_errors
                )
                
                processing_time = (datetime.utcnow() - start_time).total_seconds()
                
                result = MetadataExtractionResult(
                    extraction_id=extraction_id,
                    file_path=file_path,
                    file_category=file_category,
                    extraction_timestamp=start_time,
                    file_system_metadata=fs_metadata,
                    timestamp_data=timestamp_data,
                    geolocation_data=geolocation_data,
                    device_information=device_information,
                    content_metadata=content_metadata,
                    technical_metadata=technical_metadata,
                    security_metadata=security_metadata,
                    raw_metadata=raw_metadata,
                    extraction_errors=extraction_errors,
                    processing_time=processing_time,
                    confidence_score=confidence_score
                )
                
                await self._store_extraction_result(result, evidence_id)
                
                metadata_extraction_total.labels(
                    file_type=file_category.value
                ).inc()
                
                return result
                
            except Exception as e:
                logger.error(f"Error extracting metadata from {file_path}: {e}")
                raise

    def _determine_file_category(self, file_path: str) -> FileCategory:
        """Determine file category based on extension and MIME type"""
        
        try:
            extension = Path(file_path).suffix.lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            
            if extension in self.image_extensions or (mime_type and mime_type.startswith('image/')):
                return FileCategory.IMAGE
            elif extension in self.video_extensions or (mime_type and mime_type.startswith('video/')):
                return FileCategory.VIDEO
            elif extension in self.audio_extensions or (mime_type and mime_type.startswith('audio/')):
                return FileCategory.AUDIO
            elif extension in self.document_extensions or (mime_type and 'document' in mime_type):
                return FileCategory.DOCUMENT
            elif extension in self.archive_extensions:
                return FileCategory.ARCHIVE
            elif extension in {'.exe', '.dll', '.so', '.app'}:
                return FileCategory.EXECUTABLE
            else:
                return FileCategory.OTHER
                
        except Exception:
            return FileCategory.OTHER

    async def _extract_filesystem_metadata(self, file_path: str) -> FileSystemMetadata:
        """Extract file system metadata"""
        
        try:
            stat = os.stat(file_path)
            path_obj = Path(file_path)
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            mime_type = mime_type or "application/octet-stream"
            
            # Calculate checksums
            checksums = await self._calculate_checksums(file_path)
            
            # Get ownership information (Unix-like systems)
            try:
                import pwd
                import grp
                owner = pwd.getpwuid(stat.st_uid).pw_name
                group = grp.getgrgid(stat.st_gid).gr_name
            except (ImportError, KeyError):
                owner = str(stat.st_uid)
                group = str(stat.st_gid)
            
            # Get permissions
            permissions = oct(stat.st_mode)[-3:]
            
            return FileSystemMetadata(
                file_path=str(path_obj.absolute()),
                file_name=path_obj.name,
                file_size=stat.st_size,
                file_type=path_obj.suffix.lower(),
                mime_type=mime_type,
                permissions=permissions,
                owner=owner,
                group=group,
                inode=getattr(stat, 'st_ino', None),
                links_count=getattr(stat, 'st_nlink', 1),
                checksums=checksums
            )
            
        except Exception as e:
            logger.error(f"Error extracting filesystem metadata: {e}")
            raise

    async def _extract_timestamp_data(self, file_path: str) -> TimestampData:
        """Extract timestamp information"""
        
        try:
            stat = os.stat(file_path)
            
            # File system timestamps
            creation_time = datetime.fromtimestamp(getattr(stat, 'st_birthtime', stat.st_ctime))
            modification_time = datetime.fromtimestamp(stat.st_mtime)
            access_time = datetime.fromtimestamp(stat.st_atime)
            
            # Try to get timezone information
            try:
                import time
                timezone_info = time.tzname[0]
            except:
                timezone_info = None
            
            return TimestampData(
                creation_time=creation_time,
                modification_time=modification_time,
                access_time=access_time,
                metadata_time=None,  # Will be filled by format-specific extractors
                timezone_info=timezone_info
            )
            
        except Exception as e:
            logger.error(f"Error extracting timestamp data: {e}")
            raise

    async def _calculate_checksums(self, file_path: str) -> Dict[str, str]:
        """Calculate file checksums"""
        
        try:
            checksums = {}
            
            # Calculate MD5, SHA1, and SHA256
            hash_md5 = hashlib.md5()
            hash_sha1 = hashlib.sha1()
            hash_sha256 = hashlib.sha256()
            
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    hash_md5.update(chunk)
                    hash_sha1.update(chunk)
                    hash_sha256.update(chunk)
            
            checksums['md5'] = hash_md5.hexdigest()
            checksums['sha1'] = hash_sha1.hexdigest()
            checksums['sha256'] = hash_sha256.hexdigest()
            
            return checksums
            
        except Exception as e:
            logger.error(f"Error calculating checksums: {e}")
            return {}

    async def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract image-specific metadata"""
        
        try:
            result = {'raw': {}, 'geolocation': None, 'device': None, 'technical': None}
            
            with Image.open(file_path) as img:
                # Get basic image info
                width, height = img.size
                mode = img.mode
                format_name = img.format
                
                # Extract EXIF data
                exif_dict = img._getexif()
                if exif_dict:
                    exif_data = {}
                    gps_data = {}
                    device_data = {}
                    
                    for tag_id, value in exif_dict.items():
                        tag = TAGS.get(tag_id, tag_id)
                        
                        # Handle bytes values
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8', errors='replace')
                            except:
                                value = str(value)
                        
                        exif_data[str(tag)] = value
                        
                        # Extract device information
                        if tag == "Make":
                            device_data['make'] = str(value)
                        elif tag == "Model":
                            device_data['model'] = str(value)
                        elif tag == "Software":
                            device_data['software_version'] = str(value)
                        elif tag == "SerialNumber":
                            device_data['serial_number'] = str(value)
                        
                        # Extract GPS data
                        elif tag == "GPSInfo" and isinstance(value, dict):
                            gps_data = value
                    
                    result['raw']['exif'] = exif_data
                    
                    # Process GPS data
                    if gps_data:
                        geolocation = self._process_gps_data(gps_data)
                        if geolocation:
                            result['geolocation'] = geolocation
                    
                    # Create device information
                    if any(device_data.values()):
                        result['device'] = DeviceInformation(
                            make=device_data.get('make'),
                            model=device_data.get('model'),
                            serial_number=device_data.get('serial_number'),
                            software_version=device_data.get('software_version'),
                            hardware_info={}
                        )
                
                # Create technical metadata
                result['technical'] = TechnicalMetadata(
                    format_version=format_name,
                    compression=None,
                    color_space=mode,
                    bit_depth=None,
                    resolution=(width, height),
                    duration=None,
                    bitrate=None,
                    sample_rate=None,
                    channels=None
                )
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting image metadata: {e}")
            return {'raw': {}, 'geolocation': None, 'device': None, 'technical': None}

    async def _extract_video_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract video-specific metadata"""
        
        try:
            result = {'raw': {}, 'technical': None, 'content': None}
            
            # Use ffprobe to extract video metadata
            probe = ffmpeg.probe(file_path)
            
            # Extract format information
            format_info = probe.get('format', {})
            streams = probe.get('streams', [])
            
            # Find video stream
            video_stream = None
            audio_stream = None
            
            for stream in streams:
                if stream.get('codec_type') == 'video' and not video_stream:
                    video_stream = stream
                elif stream.get('codec_type') == 'audio' and not audio_stream:
                    audio_stream = stream
            
            # Extract technical metadata
            if video_stream:
                width = video_stream.get('width')
                height = video_stream.get('height')
                duration = float(format_info.get('duration', 0))
                bitrate = int(format_info.get('bit_rate', 0))
                
                result['technical'] = TechnicalMetadata(
                    format_version=format_info.get('format_name'),
                    compression=video_stream.get('codec_name'),
                    color_space=video_stream.get('pix_fmt'),
                    bit_depth=None,
                    resolution=(width, height) if width and height else None,
                    duration=duration if duration > 0 else None,
                    bitrate=bitrate if bitrate > 0 else None,
                    sample_rate=int(audio_stream.get('sample_rate', 0)) if audio_stream else None,
                    channels=int(audio_stream.get('channels', 0)) if audio_stream else None
                )
            
            # Extract content metadata from format tags
            tags = format_info.get('tags', {})
            if tags:
                result['content'] = ContentMetadata(
                    title=tags.get('title'),
                    author=tags.get('artist') or tags.get('author'),
                    subject=tags.get('subject'),
                    keywords=[],
                    description=tags.get('description') or tags.get('comment'),
                    creator_tool=tags.get('encoder'),
                    language=tags.get('language'),
                    page_count=None,
                    word_count=None
                )
            
            result['raw']['ffprobe'] = probe
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting video metadata: {e}")
            return {'raw': {}, 'technical': None, 'content': None}

    async def _extract_audio_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract audio-specific metadata"""
        
        try:
            result = {'raw': {}, 'technical': None, 'content': None}
            
            # Use mutagen to extract audio metadata
            audio_file = mutagen.File(file_path)
            if audio_file:
                # Extract technical information
                info = audio_file.info
                if info:
                    result['technical'] = TechnicalMetadata(
                        format_version=None,
                        compression=getattr(info, 'codec', None),
                        color_space=None,
                        bit_depth=getattr(info, 'bits_per_sample', None),
                        resolution=None,
                        duration=getattr(info, 'length', None),
                        bitrate=getattr(info, 'bitrate', None),
                        sample_rate=getattr(info, 'sample_rate', None),
                        channels=getattr(info, 'channels', None)
                    )
                
                # Extract tags
                tags = dict(audio_file.tags) if audio_file.tags else {}
                if tags:
                    # Convert tag values to strings
                    string_tags = {}
                    for key, value in tags.items():
                        if isinstance(value, list):
                            string_tags[key] = [str(v) for v in value]
                        else:
                            string_tags[key] = str(value)
                    
                    result['content'] = ContentMetadata(
                        title=string_tags.get('TIT2', [None])[0] if 'TIT2' in string_tags else None,
                        author=string_tags.get('TPE1', [None])[0] if 'TPE1' in string_tags else None,
                        subject=string_tags.get('TALB', [None])[0] if 'TALB' in string_tags else None,
                        keywords=[],
                        description=string_tags.get('COMM', [None])[0] if 'COMM' in string_tags else None,
                        creator_tool=None,
                        language=None,
                        page_count=None,
                        word_count=None
                    )
                    
                    result['raw']['tags'] = string_tags
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting audio metadata: {e}")
            return {'raw': {}, 'technical': None, 'content': None}

    async def _extract_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract document-specific metadata"""
        
        try:
            result = {'raw': {}, 'content': None, 'security': None}
            extension = Path(file_path).suffix.lower()
            
            if extension == '.pdf':
                result.update(await self._extract_pdf_metadata(file_path))
            elif extension in ['.doc', '.docx']:
                result.update(await self._extract_word_metadata(file_path))
            elif extension in ['.xls', '.xlsx']:
                result.update(await self._extract_excel_metadata(file_path))
            elif extension == '.txt':
                result.update(await self._extract_text_metadata(file_path))
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting document metadata: {e}")
            return {'raw': {}, 'content': None, 'security': None}

    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata"""
        
        try:
            result = {'raw': {}, 'content': None, 'security': None}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract document info
                if pdf_reader.metadata:
                    metadata = pdf_reader.metadata
                    
                    result['content'] = ContentMetadata(
                        title=metadata.get('/Title'),
                        author=metadata.get('/Author'),
                        subject=metadata.get('/Subject'),
                        keywords=metadata.get('/Keywords', '').split(',') if metadata.get('/Keywords') else [],
                        description=None,
                        creator_tool=metadata.get('/Creator'),
                        language=None,
                        page_count=len(pdf_reader.pages),
                        word_count=None
                    )
                    
                    result['raw']['pdf_metadata'] = {str(k): str(v) for k, v in metadata.items()}
                
                # Check security
                is_encrypted = pdf_reader.is_encrypted
                result['security'] = SecurityMetadata(
                    encryption_status=is_encrypted,
                    digital_signature=None,
                    certificate_info=None,
                    password_protected=is_encrypted,
                    access_restrictions=[]
                )
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {'raw': {}, 'content': None, 'security': None}

    async def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract Word document metadata"""
        
        try:
            result = {'raw': {}, 'content': None, 'security': None}
            
            if Path(file_path).suffix.lower() == '.docx':
                doc = docx.Document(file_path)
                
                # Extract core properties
                core_props = doc.core_properties
                
                result['content'] = ContentMetadata(
                    title=core_props.title,
                    author=core_props.author,
                    subject=core_props.subject,
                    keywords=core_props.keywords.split(',') if core_props.keywords else [],
                    description=core_props.comments,
                    creator_tool=None,
                    language=core_props.language,
                    page_count=None,
                    word_count=None
                )
                
                # Count paragraphs and estimate word count
                paragraph_count = len(doc.paragraphs)
                word_count = sum(len(p.text.split()) for p in doc.paragraphs)
                
                if result['content']:
                    result['content'].word_count = word_count
                
                result['raw']['word_metadata'] = {
                    'paragraph_count': paragraph_count,
                    'created': str(core_props.created) if core_props.created else None,
                    'modified': str(core_props.modified) if core_props.modified else None,
                    'last_modified_by': core_props.last_modified_by
                }
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting Word metadata: {e}")
            return {'raw': {}, 'content': None, 'security': None}

    async def _extract_excel_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract Excel metadata"""
        
        try:
            result = {'raw': {}, 'content': None, 'security': None}
            
            if Path(file_path).suffix.lower() == '.xlsx':
                workbook = openpyxl.load_workbook(file_path)
                
                # Extract properties
                props = workbook.properties
                
                result['content'] = ContentMetadata(
                    title=props.title,
                    author=props.creator,
                    subject=props.subject,
                    keywords=props.keywords.split(',') if props.keywords else [],
                    description=props.description,
                    creator_tool=None,
                    language=None,
                    page_count=len(workbook.worksheets),
                    word_count=None
                )
                
                # Count sheets and cells
                sheet_count = len(workbook.worksheets)
                total_cells = 0
                
                for sheet in workbook.worksheets:
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value is not None:
                                total_cells += 1
                
                result['raw']['excel_metadata'] = {
                    'sheet_count': sheet_count,
                    'total_cells_with_data': total_cells,
                    'created': str(props.created) if props.created else None,
                    'modified': str(props.modified) if props.modified else None
                }
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting Excel metadata: {e}")
            return {'raw': {}, 'content': None, 'security': None}

    async def _extract_text_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract text file metadata"""
        
        try:
            result = {'raw': {}, 'content': None, 'security': None}
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Basic text analysis
                lines = content.split('\n')
                words = content.split()
                characters = len(content)
                
                result['content'] = ContentMetadata(
                    title=None,
                    author=None,
                    subject=None,
                    keywords=[],
                    description=None,
                    creator_tool=None,
                    language=None,
                    page_count=None,
                    word_count=len(words)
                )
                
                result['raw']['text_metadata'] = {
                    'line_count': len(lines),
                    'word_count': len(words),
                    'character_count': characters,
                    'encoding': 'utf-8'
                }
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text metadata: {e}")
            return {'raw': {}, 'content': None, 'security': None}

    async def _extract_security_metadata(self, file_path: str) -> SecurityMetadata:
        """Extract security-related metadata"""
        
        try:
            # Basic security checks
            stat = os.stat(file_path)
            permissions = oct(stat.st_mode)[-3:]
            
            # Check if file is executable
            is_executable = os.access(file_path, os.X_OK)
            
            # Check file signature for encryption indicators
            encryption_indicators = await self._check_encryption_indicators(file_path)
            
            access_restrictions = []
            if permissions[0] < '6':  # Owner doesn't have full access
                access_restrictions.append("Limited owner access")
            if permissions[1] == '0':  # No group access
                access_restrictions.append("No group access")
            if permissions[2] == '0':  # No other access
                access_restrictions.append("No public access")
            
            return SecurityMetadata(
                encryption_status=encryption_indicators['encrypted'],
                digital_signature=None,  # Would require specialized analysis
                certificate_info=None,
                password_protected=encryption_indicators['password_protected'],
                access_restrictions=access_restrictions
            )
            
        except Exception as e:
            logger.error(f"Error extracting security metadata: {e}")
            return SecurityMetadata(False, None, None, False, [])

    async def _check_encryption_indicators(self, file_path: str) -> Dict[str, bool]:
        """Check for encryption indicators in file"""
        
        try:
            # Read first few bytes to check for encryption signatures
            with open(file_path, 'rb') as file:
                header = file.read(1024)
            
            # Common encryption signatures
            encryption_signatures = [
                b'Salted__',  # OpenSSL
                b'-----BEGIN PGP',  # PGP
                b'\x50\x4b\x03\x04',  # ZIP (might be encrypted)
            ]
            
            encrypted = any(sig in header for sig in encryption_signatures)
            
            # Check for password protection indicators
            password_indicators = [
                b'password',
                b'encrypted',
                b'protected'
            ]
            
            password_protected = any(ind in header.lower() for ind in password_indicators)
            
            return {
                'encrypted': encrypted,
                'password_protected': password_protected
            }
            
        except Exception:
            return {'encrypted': False, 'password_protected': False}

    def _process_gps_data(self, gps_data: Dict) -> Optional[GeolocationData]:
        """Process GPS data from EXIF"""
        
        try:
            def convert_to_degrees(value):
                d, m, s = value
                return d + (m / 60.0) + (s / 3600.0)
            
            lat = gps_data.get(2)  # GPSLatitude
            lat_ref = gps_data.get(1)  # GPSLatitudeRef
            lon = gps_data.get(4)  # GPSLongitude
            lon_ref = gps_data.get(3)  # GPSLongitudeRef
            alt = gps_data.get(6)  # GPSAltitude
            
            if lat and lon:
                latitude = convert_to_degrees(lat)
                longitude = convert_to_degrees(lon)
                
                if lat_ref == 'S':
                    latitude = -latitude
                if lon_ref == 'W':
                    longitude = -longitude
                
                altitude = float(alt) if alt else None
                
                return GeolocationData(
                    latitude=latitude,
                    longitude=longitude,
                    altitude=altitude,
                    accuracy=None,
                    timestamp=None,
                    source="exif_gps"
                )
            
            return None
            
        except Exception:
            return None

    def _calculate_confidence_score(self, fs_metadata: FileSystemMetadata,
                                  timestamp_data: TimestampData,
                                  geolocation_data: Optional[GeolocationData],
                                  device_information: Optional[DeviceInformation],
                                  content_metadata: Optional[ContentMetadata],
                                  technical_metadata: Optional[TechnicalMetadata],
                                  extraction_errors: List[str]) -> float:
        """Calculate confidence score for metadata extraction"""
        
        try:
            score = 0.0
            
            # Base score for file system metadata (always available)
            score += 0.3
            
            # Bonus for timestamp data
            if timestamp_data.creation_time:
                score += 0.1
            if timestamp_data.timezone_info:
                score += 0.05
            
            # Bonus for geolocation data
            if geolocation_data:
                score += 0.15
            
            # Bonus for device information
            if device_information and (device_information.make or device_information.model):
                score += 0.1
            
            # Bonus for content metadata
            if content_metadata:
                if content_metadata.title:
                    score += 0.05
                if content_metadata.author:
                    score += 0.05
                if content_metadata.creator_tool:
                    score += 0.05
            
            # Bonus for technical metadata
            if technical_metadata:
                score += 0.1
            
            # Penalty for extraction errors
            error_penalty = min(len(extraction_errors) * 0.1, 0.3)
            score -= error_penalty
            
            return max(0.0, min(score, 1.0))
            
        except Exception:
            return 0.5

    async def _store_extraction_result(self, result: MetadataExtractionResult, evidence_id: str = None):
        """Store metadata extraction result"""
        
        try:
            with self.Session() as session:
                record = MetadataExtractionRecord(
                    extraction_id=result.extraction_id,
                    evidence_id=evidence_id,
                    file_path=result.file_path,
                    file_category=result.file_category.value,
                    extraction_timestamp=result.extraction_timestamp,
                    file_system_metadata=asdict(result.file_system_metadata),
                    timestamp_data=asdict(result.timestamp_data),
                    geolocation_data=asdict(result.geolocation_data) if result.geolocation_data else None,
                    device_information=asdict(result.device_information) if result.device_information else None,
                    content_metadata=asdict(result.content_metadata) if result.content_metadata else None,
                    technical_metadata=asdict(result.technical_metadata) if result.technical_metadata else None,
                    security_metadata=asdict(result.security_metadata) if result.security_metadata else None,
                    raw_metadata=result.raw_metadata,
                    extraction_errors=result.extraction_errors,
                    processing_time=result.processing_time,
                    confidence_score=result.confidence_score,
                    created_at=datetime.utcnow()
                )
                
                session.add(record)
                session.commit()
                
        except Exception as e:
            logger.error(f"Error storing metadata extraction result: {e}")

def create_metadata_extractor(db_url: str = None, redis_url: str = None) -> MetadataExtractor:
    """Create and configure MetadataExtractor instance"""
    
    if not db_url:
        db_url = "postgresql://insurance_user:insurance_pass@localhost:5432/insurance_ai"
    
    if not redis_url:
        redis_url = "redis://localhost:6379/0"
    
    return MetadataExtractor(db_url=db_url, redis_url=redis_url)

