"""
Insurance AI Agent System - Document Analysis Agent
Production-ready agent for document processing, OCR, and data extraction
"""

import asyncio
import uuid
import io
import base64
from datetime import datetime
from typing import Dict, Any, List, Optional, Union, Tuple
from decimal import Decimal
import json
import re
from pathlib import Path
import tempfile
import os

# Document processing libraries
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import cv2
import numpy as np
import fitz  # PyMuPDF
from pdf2image import convert_from_bytes
import docx
from openpyxl import load_workbook

# NLP and ML libraries
import spacy
from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
import torch
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Database and utilities
from sqlalchemy.ext.asyncio import AsyncSession
import structlog
import boto3
from botocore.exceptions import ClientError

from backend.shared.models import Document, DocumentAnalysis, AgentExecution
from backend.shared.schemas import (
    DocumentCreate, DocumentAnalysisCreate,
    AgentExecutionStatus
)
from backend.shared.services import BaseService, ServiceException
from backend.shared.database import get_db_session
from backend.shared.monitoring import metrics, performance_monitor, audit_logger
from backend.shared.utils import DataUtils, ValidationUtils

logger = structlog.get_logger(__name__)

class DocumentAnalysisAgent:
    """
    Advanced document analysis agent with OCR, NLP, and intelligent data extraction
    Supports multiple document formats and provides structured data extraction
    """
    
    def __init__(self, db_session: AsyncSession):
        self.db_session = db_session
        self.agent_name = "document_analysis_agent"
        self.agent_version = "1.0.0"
        self.logger = structlog.get_logger(self.agent_name)
        
        # Initialize NLP models
        self._initialize_nlp_models()
        
        # Document type patterns
        self.document_patterns = {
            "policy_document": [
                r"policy\s+number",
                r"coverage\s+amount",
                r"premium\s+amount",
                r"effective\s+date",
                r"expiration\s+date"
            ],
            "claim_form": [
                r"claim\s+number",
                r"incident\s+date",
                r"loss\s+amount",
                r"claimant\s+name",
                r"policy\s+holder"
            ],
            "medical_report": [
                r"patient\s+name",
                r"diagnosis",
                r"treatment\s+date",
                r"physician\s+name",
                r"medical\s+history"
            ],
            "financial_statement": [
                r"balance\s+sheet",
                r"income\s+statement",
                r"cash\s+flow",
                r"assets",
                r"liabilities"
            ],
            "identity_document": [
                r"driver\s+license",
                r"passport",
                r"social\s+security",
                r"date\s+of\s+birth",
                r"address"
            ]
        }
        
        # Field extraction patterns
        self.extraction_patterns = {
            "amounts": r"\$?[\d,]+\.?\d*",
            "dates": r"\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}",
            "phone_numbers": r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}",
            "email_addresses": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}",
            "policy_numbers": r"[A-Z]{2,4}\d{6,12}",
            "claim_numbers": r"CLM\d{8,12}",
            "ssn": r"\d{3}-?\d{2}-?\d{4}",
            "zip_codes": r"\d{5}(-\d{4})?"
        }
        
        # AWS S3 client for document storage
        self.s3_client = boto3.client('s3')
        self.document_bucket = os.getenv('DOCUMENT_STORAGE_BUCKET', 'insurance-documents')
        
        # Confidence thresholds
        self.confidence_thresholds = {
            "ocr_confidence": 0.7,
            "classification_confidence": 0.8,
            "extraction_confidence": 0.75,
            "validation_confidence": 0.85
        }
    
    def _initialize_nlp_models(self):
        """Initialize NLP models and pipelines"""
        
        try:
            # Load spaCy model for NER
            self.nlp = spacy.load("en_core_web_sm")
            
            # Initialize BERT-based NER model for insurance-specific entities
            self.ner_tokenizer = AutoTokenizer.from_pretrained("dbmdz/bert-large-cased-finetuned-conll03-english")
            self.ner_model = AutoModelForTokenClassification.from_pretrained("dbmdz/bert-large-cased-finetuned-conll03-english")
            self.ner_pipeline = pipeline("ner", 
                                       model=self.ner_model, 
                                       tokenizer=self.ner_tokenizer,
                                       aggregation_strategy="simple")
            
            # Initialize sentence transformer for semantic similarity
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            
            # Initialize document classification pipeline
            self.classifier = pipeline("text-classification", 
                                     model="microsoft/DialoGPT-medium",
                                     return_all_scores=True)
            
            self.logger.info("NLP models initialized successfully")
            
        except Exception as e:
            self.logger.error("Failed to initialize NLP models", error=str(e))
            # Fallback to basic processing
            self.nlp = None
            self.ner_pipeline = None
            self.sentence_model = None
            self.classifier = None
    
    async def analyze_document(
        self,
        document_id: uuid.UUID,
        file_path: str,
        document_type: Optional[str] = None,
        analysis_options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        Analyze a document and extract structured data
        
        Args:
            document_id: UUID of the document record
            file_path: Path to the document file or S3 key
            document_type: Optional document type hint
            analysis_options: Additional analysis configuration
            
        Returns:
            Dictionary containing analysis results
        """
        
        async with performance_monitor.monitor_operation("document_analysis"):
            try:
                # Create agent execution record
                execution = AgentExecution(
                    agent_name=self.agent_name,
                    agent_version=self.agent_version,
                    input_data={
                        "document_id": str(document_id),
                        "file_path": file_path,
                        "document_type": document_type,
                        "analysis_options": analysis_options or {}
                    },
                    status=AgentExecutionStatus.RUNNING
                )
                
                self.db_session.add(execution)
                await self.db_session.commit()
                await self.db_session.refresh(execution)
                
                start_time = datetime.utcnow()
                
                # Download document if it's in S3
                local_file_path = await self._download_document(file_path)
                
                # Extract text from document
                extracted_text, ocr_confidence = await self._extract_text(local_file_path)
                
                # Classify document type
                classified_type, classification_confidence = await self._classify_document(
                    extracted_text, document_type
                )
                
                # Extract structured data
                extracted_data = await self._extract_structured_data(
                    extracted_text, classified_type
                )
                
                # Perform named entity recognition
                entities = await self._extract_entities(extracted_text)
                
                # Validate extracted data
                validation_results = await self._validate_extracted_data(
                    extracted_data, classified_type
                )
                
                # Calculate overall confidence
                overall_confidence = self._calculate_overall_confidence(
                    ocr_confidence, classification_confidence, validation_results
                )
                
                # Prepare analysis results
                analysis_result = {
                    "document_id": str(document_id),
                    "document_type": classified_type,
                    "classification_confidence": classification_confidence,
                    "ocr_confidence": ocr_confidence,
                    "overall_confidence": overall_confidence,
                    "extracted_text": extracted_text,
                    "extracted_data": extracted_data,
                    "entities": entities,
                    "validation_results": validation_results,
                    "processing_metadata": {
                        "agent_name": self.agent_name,
                        "agent_version": self.agent_version,
                        "processing_time_ms": int((datetime.utcnow() - start_time).total_seconds() * 1000),
                        "file_size_bytes": os.path.getsize(local_file_path) if os.path.exists(local_file_path) else 0,
                        "analysis_options": analysis_options or {}
                    }
                }
                
                # Save analysis results to database
                await self._save_analysis_results(document_id, analysis_result)
                
                # Update execution record
                execution.status = AgentExecutionStatus.COMPLETED
                execution.output_data = analysis_result
                execution.execution_time_ms = analysis_result["processing_metadata"]["processing_time_ms"]
                execution.completed_at = datetime.utcnow()
                
                await self.db_session.commit()
                
                # Clean up temporary file
                if local_file_path != file_path and os.path.exists(local_file_path):
                    os.unlink(local_file_path)
                
                # Record metrics
                metrics.record_agent_execution(
                    self.agent_name, 
                    execution.execution_time_ms / 1000, 
                    success=True
                )
                
                # Log successful analysis
                audit_logger.log_user_action(
                    user_id="system",
                    action="document_analysis_completed",
                    resource_type="document",
                    resource_id=str(document_id),
                    details={
                        "document_type": classified_type,
                        "confidence": overall_confidence,
                        "processing_time_ms": execution.execution_time_ms
                    }
                )
                
                self.logger.info(
                    "Document analysis completed",
                    document_id=str(document_id),
                    document_type=classified_type,
                    confidence=overall_confidence,
                    processing_time_ms=execution.execution_time_ms
                )
                
                return analysis_result
                
            except Exception as e:
                # Update execution record with error
                execution.status = AgentExecutionStatus.FAILED
                execution.error_message = str(e)
                execution.completed_at = datetime.utcnow()
                
                await self.db_session.commit()
                
                # Record metrics
                metrics.record_agent_execution(self.agent_name, 0, success=False)
                
                self.logger.error(
                    "Document analysis failed",
                    document_id=str(document_id),
                    error=str(e)
                )
                raise ServiceException(f"Document analysis failed: {str(e)}")
    
    async def _download_document(self, file_path: str) -> str:
        """Download document from S3 or return local path"""
        
        if file_path.startswith('s3://') or not file_path.startswith('/'):
            # Download from S3
            try:
                # Parse S3 path
                if file_path.startswith('s3://'):
                    bucket, key = file_path[5:].split('/', 1)
                else:
                    bucket = self.document_bucket
                    key = file_path
                
                # Create temporary file
                temp_file = tempfile.NamedTemporaryFile(delete=False)
                temp_file.close()
                
                # Download file
                self.s3_client.download_file(bucket, key, temp_file.name)
                
                return temp_file.name
                
            except ClientError as e:
                raise ServiceException(f"Failed to download document from S3: {str(e)}")
        
        else:
            # Local file path
            if not os.path.exists(file_path):
                raise ServiceException(f"Document file not found: {file_path}")
            
            return file_path
    
    async def _extract_text(self, file_path: str) -> Tuple[str, float]:
        """Extract text from document using OCR and text extraction"""
        
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return await self._extract_text_from_image(file_path)
            elif file_extension == '.docx':
                return await self._extract_text_from_docx(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return await self._extract_text_from_excel(file_path)
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read(), 1.0
            else:
                raise ServiceException(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            self.logger.error("Text extraction failed", file_path=file_path, error=str(e))
            raise ServiceException(f"Text extraction failed: {str(e)}")
    
    async def _extract_text_from_pdf(self, file_path: str) -> Tuple[str, float]:
        """Extract text from PDF using PyMuPDF and OCR fallback"""
        
        try:
            # Try text extraction first
            doc = fitz.open(file_path)
            text = ""
            total_confidence = 0.0
            page_count = 0
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                
                if page_text.strip():
                    # Text found, use direct extraction
                    text += page_text + "\n"
                    total_confidence += 1.0
                else:
                    # No text found, use OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Perform OCR
                    ocr_result = pytesseract.image_to_data(
                        image, 
                        output_type=pytesseract.Output.DICT,
                        config='--psm 6'
                    )
                    
                    # Extract text and calculate confidence
                    page_text, page_confidence = self._process_ocr_result(ocr_result)
                    text += page_text + "\n"
                    total_confidence += page_confidence
                
                page_count += 1
            
            doc.close()
            
            average_confidence = total_confidence / page_count if page_count > 0 else 0.0
            return text.strip(), average_confidence
            
        except Exception as e:
            self.logger.error("PDF text extraction failed", error=str(e))
            raise
    
    async def _extract_text_from_image(self, file_path: str) -> Tuple[str, float]:
        """Extract text from image using OCR with preprocessing"""
        
        try:
            # Load and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Enhance image for better OCR
            image = self._enhance_image_for_ocr(image)
            
            # Perform OCR
            ocr_result = pytesseract.image_to_data(
                image,
                output_type=pytesseract.Output.DICT,
                config='--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,;:!?-()[]{}/@#$%^&*+=<>|\\~`"\''
            )
            
            # Process OCR result
            text, confidence = self._process_ocr_result(ocr_result)
            
            return text, confidence
            
        except Exception as e:
            self.logger.error("Image text extraction failed", error=str(e))
            raise
    
    def _enhance_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Enhance image quality for better OCR results"""
        
        try:
            # Convert to numpy array for OpenCV processing
            img_array = np.array(image)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Apply Gaussian blur to reduce noise
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            
            # Apply adaptive thresholding
            thresh = cv2.adaptiveThreshold(
                blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Apply morphological operations to clean up
            kernel = np.ones((2, 2), np.uint8)
            cleaned = cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)
            
            # Convert back to PIL Image
            enhanced_image = Image.fromarray(cleaned)
            
            # Additional PIL enhancements
            enhanced_image = enhanced_image.filter(ImageFilter.MedianFilter())
            enhancer = ImageEnhance.Contrast(enhanced_image)
            enhanced_image = enhancer.enhance(1.5)
            
            return enhanced_image
            
        except Exception as e:
            self.logger.warning("Image enhancement failed, using original", error=str(e))
            return image
    
    def _process_ocr_result(self, ocr_result: Dict[str, List]) -> Tuple[str, float]:
        """Process OCR result and calculate confidence"""
        
        try:
            words = []
            confidences = []
            
            for i, conf in enumerate(ocr_result['conf']):
                if int(conf) > 0:  # Only include words with confidence > 0
                    word = ocr_result['text'][i].strip()
                    if word:
                        words.append(word)
                        confidences.append(int(conf))
            
            # Join words into text
            text = ' '.join(words)
            
            # Calculate average confidence
            avg_confidence = sum(confidences) / len(confidences) / 100.0 if confidences else 0.0
            
            return text, avg_confidence
            
        except Exception as e:
            self.logger.error("OCR result processing failed", error=str(e))
            return "", 0.0
    
    async def _extract_text_from_docx(self, file_path: str) -> Tuple[str, float]:
        """Extract text from DOCX file"""
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip(), 1.0  # Perfect confidence for native text extraction
            
        except Exception as e:
            self.logger.error("DOCX text extraction failed", error=str(e))
            raise
    
    async def _extract_text_from_excel(self, file_path: str) -> Tuple[str, float]:
        """Extract text from Excel file"""
        
        try:
            workbook = load_workbook(file_path, data_only=True)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                
                text += "\n"
            
            return text.strip(), 1.0  # Perfect confidence for native text extraction
            
        except Exception as e:
            self.logger.error("Excel text extraction failed", error=str(e))
            raise
    
    async def _classify_document(self, text: str, hint: Optional[str] = None) -> Tuple[str, float]:
        """Classify document type based on content"""
        
        try:
            if hint and hint in self.document_patterns:
                # Validate hint against content
                patterns = self.document_patterns[hint]
                matches = sum(1 for pattern in patterns if re.search(pattern, text, re.IGNORECASE))
                confidence = matches / len(patterns)
                
                if confidence >= self.confidence_thresholds["classification_confidence"]:
                    return hint, confidence
            
            # Classify based on content patterns
            best_type = "unknown"
            best_confidence = 0.0
            
            for doc_type, patterns in self.document_patterns.items():
                matches = sum(1 for pattern in patterns if re.search(pattern, text, re.IGNORECASE))
                confidence = matches / len(patterns)
                
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_type = doc_type
            
            # Use ML classification if available and confidence is low
            if self.classifier and best_confidence < self.confidence_thresholds["classification_confidence"]:
                try:
                    # Truncate text for classification
                    classification_text = text[:512]
                    ml_results = self.classifier(classification_text)
                    
                    # Map ML results to document types (simplified)
                    if ml_results:
                        ml_confidence = max(result['score'] for result in ml_results)
                        if ml_confidence > best_confidence:
                            best_confidence = ml_confidence
                            # Map to closest document type based on ML result
                            best_type = self._map_ml_result_to_doc_type(ml_results)
                
                except Exception as e:
                    self.logger.warning("ML classification failed", error=str(e))
            
            return best_type, best_confidence
            
        except Exception as e:
            self.logger.error("Document classification failed", error=str(e))
            return "unknown", 0.0
    
    def _map_ml_result_to_doc_type(self, ml_results: List[Dict[str, Any]]) -> str:
        """Map ML classification result to document type"""
        
        # Simplified mapping - in production, this would be more sophisticated
        highest_score_label = max(ml_results, key=lambda x: x['score'])['label']
        
        # Map common labels to document types
        label_mapping = {
            'LABEL_0': 'policy_document',
            'LABEL_1': 'claim_form',
            'LABEL_2': 'medical_report',
            'LABEL_3': 'financial_statement',
            'LABEL_4': 'identity_document'
        }
        
        return label_mapping.get(highest_score_label, 'unknown')
    
    async def _extract_structured_data(self, text: str, document_type: str) -> Dict[str, Any]:
        """Extract structured data based on document type"""
        
        try:
            extracted_data = {}
            
            # Extract common patterns
            for field_name, pattern in self.extraction_patterns.items():
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    extracted_data[field_name] = matches
            
            # Document-specific extraction
            if document_type == "policy_document":
                extracted_data.update(await self._extract_policy_data(text))
            elif document_type == "claim_form":
                extracted_data.update(await self._extract_claim_data(text))
            elif document_type == "medical_report":
                extracted_data.update(await self._extract_medical_data(text))
            elif document_type == "financial_statement":
                extracted_data.update(await self._extract_financial_data(text))
            elif document_type == "identity_document":
                extracted_data.update(await self._extract_identity_data(text))
            
            return extracted_data
            
        except Exception as e:
            self.logger.error("Structured data extraction failed", error=str(e))
            return {}
    
    async def _extract_policy_data(self, text: str) -> Dict[str, Any]:
        """Extract policy-specific data"""
        
        policy_data = {}
        
        # Policy number
        policy_match = re.search(r"policy\s+(?:number|#)?\s*:?\s*([A-Z0-9-]+)", text, re.IGNORECASE)
        if policy_match:
            policy_data["policy_number"] = policy_match.group(1)
        
        # Coverage amount
        coverage_match = re.search(r"coverage\s+amount\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if coverage_match:
            policy_data["coverage_amount"] = coverage_match.group(1).replace(',', '')
        
        # Premium amount
        premium_match = re.search(r"premium\s+amount\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if premium_match:
            policy_data["premium_amount"] = premium_match.group(1).replace(',', '')
        
        # Effective date
        effective_match = re.search(r"effective\s+date\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text, re.IGNORECASE)
        if effective_match:
            policy_data["effective_date"] = effective_match.group(1)
        
        # Expiration date
        expiration_match = re.search(r"expir(?:ation|y)\s+date\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text, re.IGNORECASE)
        if expiration_match:
            policy_data["expiration_date"] = expiration_match.group(1)
        
        return policy_data
    
    async def _extract_claim_data(self, text: str) -> Dict[str, Any]:
        """Extract claim-specific data"""
        
        claim_data = {}
        
        # Claim number
        claim_match = re.search(r"claim\s+(?:number|#)?\s*:?\s*([A-Z0-9-]+)", text, re.IGNORECASE)
        if claim_match:
            claim_data["claim_number"] = claim_match.group(1)
        
        # Incident date
        incident_match = re.search(r"incident\s+date\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text, re.IGNORECASE)
        if incident_match:
            claim_data["incident_date"] = incident_match.group(1)
        
        # Loss amount
        loss_match = re.search(r"loss\s+amount\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if loss_match:
            claim_data["loss_amount"] = loss_match.group(1).replace(',', '')
        
        # Claimant name
        claimant_match = re.search(r"claimant\s+name\s*:?\s*([A-Za-z\s]+)", text, re.IGNORECASE)
        if claimant_match:
            claim_data["claimant_name"] = claimant_match.group(1).strip()
        
        return claim_data
    
    async def _extract_medical_data(self, text: str) -> Dict[str, Any]:
        """Extract medical report data"""
        
        medical_data = {}
        
        # Patient name
        patient_match = re.search(r"patient\s+name\s*:?\s*([A-Za-z\s]+)", text, re.IGNORECASE)
        if patient_match:
            medical_data["patient_name"] = patient_match.group(1).strip()
        
        # Diagnosis
        diagnosis_match = re.search(r"diagnosis\s*:?\s*([A-Za-z\s,.-]+)", text, re.IGNORECASE)
        if diagnosis_match:
            medical_data["diagnosis"] = diagnosis_match.group(1).strip()
        
        # Treatment date
        treatment_match = re.search(r"treatment\s+date\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text, re.IGNORECASE)
        if treatment_match:
            medical_data["treatment_date"] = treatment_match.group(1)
        
        return medical_data
    
    async def _extract_financial_data(self, text: str) -> Dict[str, Any]:
        """Extract financial statement data"""
        
        financial_data = {}
        
        # Total assets
        assets_match = re.search(r"total\s+assets\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if assets_match:
            financial_data["total_assets"] = assets_match.group(1).replace(',', '')
        
        # Total liabilities
        liabilities_match = re.search(r"total\s+liabilities\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if liabilities_match:
            financial_data["total_liabilities"] = liabilities_match.group(1).replace(',', '')
        
        # Net income
        income_match = re.search(r"net\s+income\s*:?\s*\$?([\d,]+\.?\d*)", text, re.IGNORECASE)
        if income_match:
            financial_data["net_income"] = income_match.group(1).replace(',', '')
        
        return financial_data
    
    async def _extract_identity_data(self, text: str) -> Dict[str, Any]:
        """Extract identity document data"""
        
        identity_data = {}
        
        # Name
        name_match = re.search(r"(?:name|full\s+name)\s*:?\s*([A-Za-z\s]+)", text, re.IGNORECASE)
        if name_match:
            identity_data["name"] = name_match.group(1).strip()
        
        # Date of birth
        dob_match = re.search(r"(?:date\s+of\s+birth|dob)\s*:?\s*(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})", text, re.IGNORECASE)
        if dob_match:
            identity_data["date_of_birth"] = dob_match.group(1)
        
        # Address
        address_match = re.search(r"address\s*:?\s*([A-Za-z0-9\s,.-]+)", text, re.IGNORECASE)
        if address_match:
            identity_data["address"] = address_match.group(1).strip()
        
        return identity_data
    
    async def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities using NLP"""
        
        try:
            entities = []
            
            # Use spaCy for basic NER
            if self.nlp:
                doc = self.nlp(text[:1000000])  # Limit text length
                for ent in doc.ents:
                    entities.append({
                        "text": ent.text,
                        "label": ent.label_,
                        "start": ent.start_char,
                        "end": ent.end_char,
                        "confidence": 0.8  # Default confidence for spaCy
                    })
            
            # Use BERT-based NER for additional entities
            if self.ner_pipeline:
                try:
                    # Process text in chunks due to token limits
                    chunk_size = 500
                    text_chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
                    
                    for chunk in text_chunks[:5]:  # Limit to first 5 chunks
                        ner_results = self.ner_pipeline(chunk)
                        for result in ner_results:
                            entities.append({
                                "text": result["word"],
                                "label": result["entity_group"],
                                "start": result["start"],
                                "end": result["end"],
                                "confidence": result["score"]
                            })
                
                except Exception as e:
                    self.logger.warning("BERT NER failed", error=str(e))
            
            # Deduplicate entities
            unique_entities = []
            seen_entities = set()
            
            for entity in entities:
                entity_key = (entity["text"].lower(), entity["label"])
                if entity_key not in seen_entities:
                    seen_entities.add(entity_key)
                    unique_entities.append(entity)
            
            return unique_entities
            
        except Exception as e:
            self.logger.error("Entity extraction failed", error=str(e))
            return []
    
    async def _validate_extracted_data(self, extracted_data: Dict[str, Any], document_type: str) -> Dict[str, Any]:
        """Validate extracted data for consistency and accuracy"""
        
        validation_results = {
            "is_valid": True,
            "validation_score": 1.0,
            "errors": [],
            "warnings": []
        }
        
        try:
            # Validate amounts
            if "amounts" in extracted_data:
                for amount in extracted_data["amounts"]:
                    if not ValidationUtils.validate_amount(amount.replace('$', '').replace(',', '')):
                        validation_results["errors"].append(f"Invalid amount format: {amount}")
                        validation_results["is_valid"] = False
            
            # Validate dates
            if "dates" in extracted_data:
                for date_str in extracted_data["dates"]:
                    if not ValidationUtils.validate_date(date_str):
                        validation_results["errors"].append(f"Invalid date format: {date_str}")
                        validation_results["is_valid"] = False
            
            # Validate email addresses
            if "email_addresses" in extracted_data:
                for email in extracted_data["email_addresses"]:
                    if not ValidationUtils.validate_email(email):
                        validation_results["errors"].append(f"Invalid email format: {email}")
                        validation_results["is_valid"] = False
            
            # Document-specific validation
            if document_type == "policy_document":
                validation_results.update(await self._validate_policy_data(extracted_data))
            elif document_type == "claim_form":
                validation_results.update(await self._validate_claim_data(extracted_data))
            
            # Calculate validation score
            total_checks = len(extracted_data)
            error_count = len(validation_results["errors"])
            warning_count = len(validation_results["warnings"])
            
            if total_checks > 0:
                validation_results["validation_score"] = max(0.0, 1.0 - (error_count * 0.2 + warning_count * 0.1))
            
            return validation_results
            
        except Exception as e:
            self.logger.error("Data validation failed", error=str(e))
            return {
                "is_valid": False,
                "validation_score": 0.0,
                "errors": [f"Validation failed: {str(e)}"],
                "warnings": []
            }
    
    async def _validate_policy_data(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate policy-specific data"""
        
        validation_update = {"errors": [], "warnings": []}
        
        # Check if policy number exists
        if "policy_number" not in extracted_data:
            validation_update["warnings"].append("Policy number not found")
        
        # Validate coverage amount
        if "coverage_amount" in extracted_data:
            try:
                amount = float(extracted_data["coverage_amount"])
                if amount <= 0:
                    validation_update["errors"].append("Coverage amount must be positive")
            except ValueError:
                validation_update["errors"].append("Invalid coverage amount format")
        
        return validation_update
    
    async def _validate_claim_data(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate claim-specific data"""
        
        validation_update = {"errors": [], "warnings": []}
        
        # Check if claim number exists
        if "claim_number" not in extracted_data:
            validation_update["warnings"].append("Claim number not found")
        
        # Validate loss amount
        if "loss_amount" in extracted_data:
            try:
                amount = float(extracted_data["loss_amount"])
                if amount <= 0:
                    validation_update["errors"].append("Loss amount must be positive")
            except ValueError:
                validation_update["errors"].append("Invalid loss amount format")
        
        return validation_update
    
    def _calculate_overall_confidence(
        self, 
        ocr_confidence: float, 
        classification_confidence: float, 
        validation_results: Dict[str, Any]
    ) -> float:
        """Calculate overall confidence score"""
        
        try:
            # Weighted average of different confidence scores
            weights = {
                "ocr": 0.3,
                "classification": 0.3,
                "validation": 0.4
            }
            
            validation_confidence = validation_results.get("validation_score", 0.0)
            
            overall_confidence = (
                ocr_confidence * weights["ocr"] +
                classification_confidence * weights["classification"] +
                validation_confidence * weights["validation"]
            )
            
            return min(1.0, max(0.0, overall_confidence))
            
        except Exception as e:
            self.logger.error("Confidence calculation failed", error=str(e))
            return 0.0
    
    async def _save_analysis_results(self, document_id: uuid.UUID, analysis_result: Dict[str, Any]):
        """Save analysis results to database"""
        
        try:
            analysis_create = DocumentAnalysisCreate(
                document_id=document_id,
                analysis_type="full_analysis",
                analysis_data=analysis_result,
                confidence_score=Decimal(str(analysis_result["overall_confidence"])),
                extracted_text=analysis_result["extracted_text"],
                entities=analysis_result["entities"],
                validation_results=analysis_result["validation_results"]
            )
            
            analysis_service = BaseService(DocumentAnalysis, self.db_session)
            await analysis_service.create(analysis_create)
            
        except Exception as e:
            self.logger.error("Failed to save analysis results", error=str(e))
            # Don't raise exception as this is not critical for the analysis itself
    
    async def batch_analyze_documents(
        self,
        document_ids: List[uuid.UUID],
        analysis_options: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """Analyze multiple documents in batch"""
        
        try:
            results = []
            
            # Process documents in parallel (with concurrency limit)
            semaphore = asyncio.Semaphore(5)  # Limit to 5 concurrent analyses
            
            async def analyze_single(doc_id):
                async with semaphore:
                    try:
                        # Get document info from database
                        doc_service = BaseService(Document, self.db_session)
                        document = await doc_service.get(doc_id)
                        
                        if not document:
                            return {
                                "document_id": str(doc_id),
                                "error": "Document not found"
                            }
                        
                        # Analyze document
                        result = await self.analyze_document(
                            doc_id,
                            document.file_path,
                            document.document_type,
                            analysis_options
                        )
                        
                        return result
                        
                    except Exception as e:
                        return {
                            "document_id": str(doc_id),
                            "error": str(e)
                        }
            
            # Execute all analyses
            tasks = [analyze_single(doc_id) for doc_id in document_ids]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results
            processed_results = []
            for result in results:
                if isinstance(result, Exception):
                    processed_results.append({
                        "error": str(result)
                    })
                else:
                    processed_results.append(result)
            
            self.logger.info(
                "Batch document analysis completed",
                total_documents=len(document_ids),
                successful_analyses=len([r for r in processed_results if "error" not in r])
            )
            
            return processed_results
            
        except Exception as e:
            self.logger.error("Batch document analysis failed", error=str(e))
            raise ServiceException(f"Batch analysis failed: {str(e)}")

# Agent factory function
async def create_document_analysis_agent(db_session: AsyncSession) -> DocumentAnalysisAgent:
    """Create document analysis agent instance"""
    return DocumentAnalysisAgent(db_session)

